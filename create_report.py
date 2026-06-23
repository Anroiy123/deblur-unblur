"""Inspect formatting metadata from a reference DOCX document."""

import argparse
from pathlib import Path


def build_parser():
    parser = argparse.ArgumentParser(
        description="Inspect structure and formatting of a reference DOCX file."
    )
    parser.add_argument("reference", type=Path, help="Path to the reference DOCX file.")
    parser.add_argument(
        "--output",
        type=Path,
        default=Path.cwd() / "BaoCao_DeBlur_analysis.txt",
        help="Text file that receives the analysis.",
    )
    return parser


def parse_args(argv=None):
    return build_parser().parse_args(argv)


def inspect_document(reference_path):
    try:
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError as exc:
        raise RuntimeError(
            "Missing python-docx. Install it with: python -m pip install python-docx"
        ) from exc

    reference_path = Path(reference_path)
    if not reference_path.is_file():
        raise FileNotFoundError(f"Reference file not found: {reference_path}")

    document = Document(reference_path)
    lines = [f"Reference: {reference_path.resolve()}", ""]

    lines.append("=== PARAGRAPH STYLES ===")
    for style in document.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH or not style.name:
            continue
        font = style.font
        paragraph_format = style.paragraph_format
        lines.append(
            f"{style.name}: font={font.name}, size={font.size}, "
            f"bold={font.bold}, italic={font.italic}, "
            f"alignment={paragraph_format.alignment}"
        )

    lines.extend(["", "=== DOCUMENT SECTIONS ==="])
    for index, section in enumerate(document.sections):
        lines.append(
            f"Section {index}: page={section.page_width}x{section.page_height}, "
            f"margins=({section.left_margin}, {section.right_margin}, "
            f"{section.top_margin}, {section.bottom_margin})"
        )

    lines.extend(["", "=== FIRST 30 PARAGRAPHS ==="])
    for index, paragraph in enumerate(document.paragraphs[:30]):
        style_name = paragraph.style.name if paragraph.style else "None"
        lines.append(
            f"[{index}] style={style_name!r}, alignment={paragraph.alignment}: "
            f"{paragraph.text[:120]!r}"
        )

    lines.extend(["", "=== TABLES ==="])
    for index, table in enumerate(document.tables):
        lines.append(
            f"Table {index}: {len(table.rows)} rows x {len(table.columns)} columns"
        )

    image_count = sum(
        1 for relation in document.part.rels.values() if "image" in relation.reltype
    )
    lines.extend(["", f"Images: {image_count}"])
    return "\n".join(lines) + "\n"


def main(argv=None):
    args = parse_args(argv)
    report = inspect_document(args.reference)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(report, encoding="utf-8")
    print(f"Analysis written to: {args.output.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
